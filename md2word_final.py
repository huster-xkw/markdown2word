import re
import os
from copy import deepcopy
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
import subprocess
import tempfile


LIST_ITEM_RE = re.compile(
    r"^(?P<indent>\s*)(?P<marker>(?:[-+*]|\d+[.)]))\s+(?P<content>.*)$"
)
HORIZONTAL_RULE_RE = re.compile(
    r"^(?:(?:-{3,})|(?:\*{3,})|(?:_{3,})|(?:-\s*){3,}|(?:\*\s*){3,}|(?:_\s*){3,})$"
)
FENCE_START_RE = re.compile(r"^(?P<fence>`{3,}|~{3,})(?P<lang>.*)$")
BARE_URL_RE = re.compile(r"https?://[^\s<>()]+(?:\([^\s<>()]*\)[^\s<>()]*)*")
LINK_LABEL_LINE_RE = re.compile(
    r"^\s*(?P<label>[^:：\n]{1,24})[：:]\s*(?P<link>(?:\[[^\]]+\]\([^)]+\))|(?:https?://\S+))\s*$",
    re.IGNORECASE,
)

FORMULA_NUMBERING_NONE = "none"
FORMULA_NUMBERING_CHAPTER_INDEX = "chapter_index"
FORMULA_NUMBERING_GLOBAL = "global"
FORMULA_NUMBERING_CHAPTER = "chapter"

FORMULA_NUMBERING_MODE_ALIASES = {
    "none": FORMULA_NUMBERING_NONE,
    "off": FORMULA_NUMBERING_NONE,
    "no": FORMULA_NUMBERING_NONE,
    "不编号": FORMULA_NUMBERING_NONE,
    "(章-序号)": FORMULA_NUMBERING_CHAPTER_INDEX,
    "章-序号": FORMULA_NUMBERING_CHAPTER_INDEX,
    "chapter_index": FORMULA_NUMBERING_CHAPTER_INDEX,
    "chapter-index": FORMULA_NUMBERING_CHAPTER_INDEX,
    "global": FORMULA_NUMBERING_GLOBAL,
    "全文连续编号": FORMULA_NUMBERING_GLOBAL,
    "全文连续": FORMULA_NUMBERING_GLOBAL,
    "chapter": FORMULA_NUMBERING_CHAPTER,
    "章节连续编号": FORMULA_NUMBERING_CHAPTER,
    "章节连续": FORMULA_NUMBERING_CHAPTER,
}

DEFAULT_DOC_STYLE_OPTIONS = {
    "font_size_body": 11,
    "font_size_h1": 16,
    "font_size_h2": 14,
    "font_size_h3": 13,
    "font_size_h4": 12,
    "font_zh": "微软雅黑",
    "font_en": "Calibri",
    "line_spacing": 1.5,
    "paragraph_spacing": 6.0,
}


def _to_int(value, default):
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return default


def _to_float(value, default):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def normalize_formula_numbering_mode(mode, enable_block_formula_numbering=True):
    if mode is None:
        return (
            FORMULA_NUMBERING_CHAPTER_INDEX
            if enable_block_formula_numbering
            else FORMULA_NUMBERING_NONE
        )

    normalized = str(mode).strip().lower()
    if not normalized:
        return FORMULA_NUMBERING_CHAPTER_INDEX
    return FORMULA_NUMBERING_MODE_ALIASES.get(normalized, FORMULA_NUMBERING_CHAPTER_INDEX)


def normalize_doc_style_options(style_options=None):
    options = dict(DEFAULT_DOC_STYLE_OPTIONS)
    if not isinstance(style_options, dict):
        return options

    options["font_zh"] = str(style_options.get("font_zh", options["font_zh"])).strip() or options["font_zh"]
    options["font_en"] = str(style_options.get("font_en", options["font_en"])).strip() or options["font_en"]

    options["font_size_body"] = max(8, min(32, _to_int(style_options.get("font_size_body"), options["font_size_body"])))
    options["font_size_h1"] = max(8, min(48, _to_int(style_options.get("font_size_h1"), options["font_size_h1"])))
    options["font_size_h2"] = max(8, min(48, _to_int(style_options.get("font_size_h2"), options["font_size_h2"])))
    options["font_size_h3"] = max(8, min(48, _to_int(style_options.get("font_size_h3"), options["font_size_h3"])))
    options["font_size_h4"] = max(8, min(48, _to_int(style_options.get("font_size_h4"), options["font_size_h4"])))

    options["line_spacing"] = max(1.0, min(3.0, _to_float(style_options.get("line_spacing"), options["line_spacing"])))
    options["paragraph_spacing"] = max(0.0, min(48.0, _to_float(style_options.get("paragraph_spacing"), options["paragraph_spacing"])))
    return options

def parse_html_table(table_html):
    """
    解析 HTML 表格，提取表格结构
    """
    soup = BeautifulSoup(table_html, 'html.parser')
    table = soup.find('table')

    if not table:
        return None

    rows = []
    for tr in table.find_all('tr'):
        cells = []
        for cell in tr.find_all(['td', 'th']):
            cells.append({
                'text': cell.get_text(strip=True),
                'rowspan': int(cell.get('rowspan', 1)),
                'colspan': int(cell.get('colspan', 1)),
                'is_header': cell.name == 'th'
            })
        rows.append(cells)

    return rows

def calculate_table_dimensions(rows_data):
    """
    计算表格的实际行列数
    """
    if not rows_data:
        return 0, 0

    max_cols = 0
    for row in rows_data:
        col_count = sum(cell['colspan'] for cell in row)
        max_cols = max(max_cols, col_count)

    total_rows = len(rows_data)
    return total_rows, max_cols

def build_cell_matrix(rows_data, total_rows, total_cols):
    """
    构建单元格矩阵，标记哪些位置被合并占据
    """
    matrix = [[None for _ in range(total_cols)] for _ in range(total_rows)]

    for row_idx, row_cells in enumerate(rows_data):
        col_idx = 0
        for cell_info in row_cells:
            # 找到下一个空位置
            while col_idx < total_cols and matrix[row_idx][col_idx] is not None:
                col_idx += 1

            if col_idx >= total_cols:
                break

            # 标记这个单元格及其合并的区域
            for r in range(cell_info['rowspan']):
                for c in range(cell_info['colspan']):
                    if row_idx + r < total_rows and col_idx + c < total_cols:
                        matrix[row_idx + r][col_idx + c] = {
                            'source_row': row_idx,
                            'source_col': col_idx,
                            'cell_info': cell_info,
                            'is_origin': (r == 0 and c == 0)
                        }

            col_idx += cell_info['colspan']

    return matrix

def create_word_table(doc, table_html, style_options=None):
    """
    在 Word 文档中创建表格（pan4 的方法）
    """
    rows_data = parse_html_table(table_html)
    if not rows_data:
        return None

    total_rows, total_cols = calculate_table_dimensions(rows_data)
    if total_rows == 0 or total_cols == 0:
        return None

    # 创建表格
    table = doc.add_table(rows=total_rows, cols=total_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass

    # 构建单元格矩阵
    matrix = build_cell_matrix(rows_data, total_rows, total_cols)
    merged = set()

    # 第一遍：填充所有原始单元格的文本
    for row_idx in range(total_rows):
        for col_idx in range(total_cols):
            cell_data = matrix[row_idx][col_idx]
            if cell_data and cell_data['is_origin']:
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_data['cell_info']['text']

                # 设置文本格式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _apply_east_asia_font(paragraph, style_options=style_options)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # 第二遍：处理单元格合并
    for row_idx in range(total_rows):
        for col_idx in range(total_cols):
            cell_data = matrix[row_idx][col_idx]
            if cell_data and cell_data['is_origin']:
                cell_info = cell_data['cell_info']

                # 需要合并的情况
                if cell_info['rowspan'] > 1 or cell_info['colspan'] > 1:
                    end_row = min(row_idx + cell_info['rowspan'] - 1, total_rows - 1)
                    end_col = min(col_idx + cell_info['colspan'] - 1, total_cols - 1)

                    try:
                        start_cell = table.cell(row_idx, col_idx)
                        end_cell = table.cell(end_row, end_col)

                        if (row_idx, col_idx) != (end_row, end_col):
                            merge_key = (row_idx, col_idx, end_row, end_col)
                            if merge_key not in merged:
                                start_cell.merge(end_cell)
                                merged.add(merge_key)
                    except Exception as e:
                        print(f"  警告：合并单元格失败 ({row_idx},{col_idx})->({end_row},{end_col}): {e}")

    _apply_table_layout(table, apply_run_size=True)
    return table


def _find_unescaped(text, token, start_idx):
    """查找未被反斜杠转义的 token 位置。"""
    idx = start_idx
    while True:
        pos = text.find(token, idx)
        if pos == -1:
            return -1

        backslashes = 0
        cursor = pos - 1
        while cursor >= 0 and text[cursor] == "\\":
            backslashes += 1
            cursor -= 1

        if backslashes % 2 == 0:
            return pos

        idx = pos + 1


def _append_styled_run(paragraph, text, states):
    """按当前样式状态向段落追加 run。"""
    if not text:
        return

    run = paragraph.add_run(text)
    run.bold = states["bold"]
    run.italic = states["italic"]
    run.font.strike = states["strike"]
    run.font.subscript = states["subscript"]
    run.font.superscript = states["superscript"]
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW if states["highlight"] else None


def _append_code_run(paragraph, text, states):
    """按代码样式向段落追加 run。"""
    if not text:
        return

    from docx.oxml.shared import qn

    run = paragraph.add_run(text)
    run.bold = states["bold"]
    run.italic = states["italic"]
    run.font.strike = states["strike"]
    run.font.subscript = states["subscript"]
    run.font.superscript = states["superscript"]
    run.font.name = "Consolas"
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    run_pr = run._element.get_or_add_rPr()
    run_pr.rFonts.set(qn("w:eastAsia"), "Consolas")


def _normalize_link_target(raw_target):
    target = raw_target.strip()
    if target.startswith("<") and target.endswith(">"):
        target = target[1:-1].strip()
    if " " in target:
        target = target.split(" ", 1)[0]
    return target


def _is_supported_link_target(target):
    return target.startswith("http://") or target.startswith("https://")


def _simplify_link_label_line(text):
    """
    将“链接标签: 链接”格式收敛为仅保留链接本体，减少视觉冗余。
    仅处理短标签且标签语义包含链接关键词的场景。
    """
    match = LINK_LABEL_LINE_RE.match(text)
    if not match:
        return text

    label = match.group("label").strip().lower()
    if not any(keyword in label for keyword in ("链接", "link", "url", "网址")):
        return text

    return match.group("link").strip()


def _compose_hyperlink_display_text(link_text, url):
    """统一超链接显示文案：Markdown 链接显示 text (url)，裸链接保持 url。"""
    normalized_text = (link_text or "").strip()
    normalized_url = (url or "").strip()
    if not normalized_url:
        return normalized_text
    if not normalized_text or normalized_text == normalized_url:
        return normalized_url
    return f"{normalized_text} ({normalized_url})"


def _append_hyperlink_run(paragraph, url, text, states):
    """向段落追加可点击超链接。"""
    from docx.oxml import OxmlElement
    from docx.oxml.shared import qn

    if not text:
        return

    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    # 为不同 Office 客户端提供更稳定的可视化表现（蓝色+下划线）。
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    if states["bold"]:
        r_pr.append(OxmlElement("w:b"))
    if states["italic"]:
        r_pr.append(OxmlElement("w:i"))
    if states["strike"]:
        r_pr.append(OxmlElement("w:strike"))
    if states["subscript"]:
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "subscript")
        r_pr.append(vert)
    elif states["superscript"]:
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "superscript")
        r_pr.append(vert)
    if states["highlight"]:
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), "yellow")
        r_pr.append(highlight)

    run.append(r_pr)
    text_elem = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _try_toggle_token(text, idx, token, state_name, states, buffer, paragraph):
    """尝试切换样式 token；返回是否处理成功。"""
    token_len = len(token)

    if states[state_name]:
        _append_styled_run(paragraph, "".join(buffer), states)
        buffer.clear()
        states[state_name] = False
        return token_len, True

    if _find_unescaped(text, token, idx + token_len) != -1:
        _append_styled_run(paragraph, "".join(buffer), states)
        buffer.clear()
        states[state_name] = True
        return token_len, True

    return 0, False


def add_markdown_inline_runs(paragraph, text):
    """
    解析并写入常用 Markdown 行内语法：
    **加粗** *斜体* ~~删除线~~ ==高亮== ~下标~ ^上标^
    """
    states = {
        "bold": False,
        "italic": False,
        "strike": False,
        "highlight": False,
        "subscript": False,
        "superscript": False,
    }

    buffer = []
    idx = 0
    text_len = len(text)

    while idx < text_len:
        if text[idx] == "\\" and idx + 1 < text_len:
            buffer.append(text[idx + 1])
            idx += 2
            continue

        if text[idx] == "`":
            end_idx = _find_unescaped(text, "`", idx + 1)
            if end_idx != -1:
                _append_styled_run(paragraph, "".join(buffer), states)
                buffer.clear()
                _append_code_run(paragraph, text[idx + 1:end_idx], states)
                idx = end_idx + 1
                continue

        # Markdown 链接: [text](url)
        if text[idx] == "[":
            close_bracket = _find_unescaped(text, "]", idx + 1)
            if (
                close_bracket != -1
                and close_bracket + 1 < text_len
                and text[close_bracket + 1] == "("
            ):
                close_paren = _find_unescaped(text, ")", close_bracket + 2)
                if close_paren != -1:
                    link_text = text[idx + 1:close_bracket]
                    raw_target = text[close_bracket + 2:close_paren]
                    target = _normalize_link_target(raw_target)
                    if _is_supported_link_target(target):
                        _append_styled_run(paragraph, "".join(buffer), states)
                        buffer.clear()
                        display_text = _compose_hyperlink_display_text(link_text, target)
                        _append_hyperlink_run(paragraph, target, display_text, states)
                        idx = close_paren + 1
                        continue

        # 裸链接: https://...
        bare_match = BARE_URL_RE.match(text, idx)
        if bare_match:
            raw_url = bare_match.group(0)
            url = raw_url
            trailing = ""

            while url and url[-1] in ".,;:!?":
                trailing = url[-1] + trailing
                url = url[:-1]

            while url.endswith(")") and url.count("(") < url.count(")"):
                trailing = ")" + trailing
                url = url[:-1]

            if _is_supported_link_target(url):
                _append_styled_run(paragraph, "".join(buffer), states)
                buffer.clear()
                _append_hyperlink_run(paragraph, url, url, states)
                if trailing:
                    buffer.append(trailing)
                idx += len(raw_url)
                continue

        handled = False

        for token, state_name in (("**", "bold"), ("~~", "strike"), ("==", "highlight")):
            if text.startswith(token, idx):
                consumed, handled = _try_toggle_token(
                    text, idx, token, state_name, states, buffer, paragraph
                )
                if handled:
                    idx += consumed
                    break

        if handled:
            continue

        if text[idx] == "*":
            consumed, handled = _try_toggle_token(
                text, idx, "*", "italic", states, buffer, paragraph
            )
            if handled:
                idx += consumed
                continue

        if text[idx] == "~":
            consumed, handled = _try_toggle_token(
                text, idx, "~", "subscript", states, buffer, paragraph
            )
            if handled:
                idx += consumed
                continue

        if text[idx] == "^":
            consumed, handled = _try_toggle_token(
                text, idx, "^", "superscript", states, buffer, paragraph
            )
            if handled:
                idx += consumed
                continue

        buffer.append(text[idx])
        idx += 1

    _append_styled_run(paragraph, "".join(buffer), states)


def _apply_paragraph_spacing(paragraph, style_options=None):
    """应用段落行距与段后间距。"""
    if not style_options:
        return

    paragraph.paragraph_format.line_spacing = style_options["line_spacing"]
    paragraph.paragraph_format.space_after = Pt(style_options["paragraph_spacing"])


def _apply_east_asia_font(
    paragraph,
    font_name="微软雅黑",
    style_options=None,
    apply_spacing=False,
):
    """统一设置段落内 run 的中英文字体。"""
    from docx.oxml.shared import qn

    if style_options:
        zh_font = style_options["font_zh"]
        en_font = style_options["font_en"]
    else:
        zh_font = font_name
        en_font = font_name

    for run in paragraph.runs:
        run_style_val = None
        if run._element.rPr is not None and run._element.rPr.rStyle is not None:
            run_style_val = run._element.rPr.rStyle.val

        if run.font.name == "Consolas" or run_style_val == "VerbatimChar":
            run.font.name = "Consolas"
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            run_pr = run._element.get_or_add_rPr()
            run_pr.rFonts.set(qn("w:ascii"), "Consolas")
            run_pr.rFonts.set(qn("w:hAnsi"), "Consolas")
            run_pr.rFonts.set(qn("w:eastAsia"), "Consolas")
            continue

        run.font.name = en_font
        run_pr = run._element.get_or_add_rPr()
        run_pr.rFonts.set(qn("w:ascii"), en_font)
        run_pr.rFonts.set(qn("w:hAnsi"), en_font)
        run_pr.rFonts.set(qn("w:eastAsia"), zh_font)

    if apply_spacing:
        _apply_paragraph_spacing(paragraph, style_options=style_options)


def _apply_table_layout(table, apply_run_size=False):
    """统一设置表格：整体居中、单元格内容水平+垂直居中。"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if apply_run_size:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


def _remove_table_borders(table):
    """移除表格可见边框。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    # 先删旧的 tblBorders
    for child in list(tbl_pr):
        if child.tag.endswith('}tblBorders'):
            tbl_pr.remove(child)

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'nil')
        borders.append(elem)
    tbl_pr.append(borders)


def _copy_pandoc_paragraph_children(
    target_para,
    source_para_element,
    hyperlink_map=None,
):
    """复制 Pandoc 生成段落的内容到目标段落，跳过段落属性。"""
    from docx.oxml.shared import qn

    for child in source_para_element:
        if child.tag.endswith('}pPr'):
            continue

        if child.tag.endswith('}hyperlink'):
            rel_id = child.get(qn('r:id'))
            link_target = hyperlink_map.get(rel_id) if hyperlink_map else None
            link_text = ''.join(
                node.text for node in child.iter() if node.tag.endswith('}t') and node.text
            )
            if link_target and link_text:
                plain_state = {
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "highlight": False,
                    "subscript": False,
                    "superscript": False,
                }
                display_text = _compose_hyperlink_display_text(link_text, link_target)
                _append_hyperlink_run(target_para, link_target, display_text, plain_state)
                continue

        target_para._element.append(deepcopy(child))


def _extract_pandoc_paragraphs(elements):
    return [element for element in elements if element.tag.endswith('}p')]


def _update_chapter_state_by_heading(level, render_state):
    if not render_state:
        return
    mode = render_state.get("formula_numbering_mode", FORMULA_NUMBERING_CHAPTER_INDEX)
    if mode not in (FORMULA_NUMBERING_CHAPTER_INDEX, FORMULA_NUMBERING_CHAPTER):
        return

    chapter_level = render_state.get('chapter_level')
    if chapter_level is None:
        render_state['chapter_level'] = level
        chapter_level = level

    if level == chapter_level:
        render_state['current_chapter'] += 1
        render_state['chapter_formula_counter'] = 0


def _next_block_formula_label(render_state):
    mode = render_state.get("formula_numbering_mode", FORMULA_NUMBERING_CHAPTER_INDEX)

    if mode == FORMULA_NUMBERING_NONE:
        return None

    if mode == FORMULA_NUMBERING_GLOBAL:
        render_state["global_formula_counter"] += 1
        return f"({render_state['global_formula_counter']})"

    if render_state['current_chapter'] <= 0:
        render_state['current_chapter'] = 1

    render_state['chapter_formula_counter'] += 1

    if mode == FORMULA_NUMBERING_CHAPTER:
        return f"({render_state['chapter_formula_counter']})"

    return f"({render_state['current_chapter']}-{render_state['chapter_formula_counter']})"


def _append_numbered_formula_block(
    doc,
    elements,
    label,
    hyperlink_map=None,
    style_options=None,
):
    """
    渲染“公式 + 右侧编号”同一行：
    使用无边框两列表格，左侧放公式，右侧放编号。
    """
    formula_paragraphs = _extract_pandoc_paragraphs(elements)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    formula_cell = table.cell(0, 0)
    label_cell = table.cell(0, 1)
    formula_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    formula_paras = []
    if formula_paragraphs:
        base_para = formula_cell.paragraphs[0]
        for idx, src_para in enumerate(formula_paragraphs):
            target_para = base_para if idx == 0 else formula_cell.add_paragraph()
            _copy_pandoc_paragraph_children(
                target_para,
                src_para,
                hyperlink_map=hyperlink_map,
            )
            target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _apply_east_asia_font(
                target_para,
                style_options=style_options,
                apply_spacing=True,
            )
            formula_paras.append(target_para)
    else:
        fallback_para = formula_cell.paragraphs[0]
        fallback_para.add_run("")
        fallback_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_east_asia_font(
            fallback_para,
            style_options=style_options,
            apply_spacing=True,
        )
        formula_paras.append(fallback_para)

    label_para = label_cell.paragraphs[0]
    label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label_para.add_run(label)
    _apply_east_asia_font(
        label_para,
        style_options=style_options,
        apply_spacing=True,
    )

    # 列宽比例：公式区域更宽，编号区域收窄
    formula_cell.width = Inches(5.8)
    label_cell.width = Inches(1.0)


def _append_pandoc_inline_to_paragraph(paragraph, text):
    """
    将 Pandoc 解析后的段落级内容（不含段落属性）追加到现有段落。
    主要用于列表项中含公式的场景，避免覆盖列表样式。
    """
    payload = convert_text_with_pandoc(text)
    elements = payload["elements"]
    hyperlink_map = payload["hyperlinks"]
    for element in elements:
        if not element.tag.endswith("}p"):
            continue

        _copy_pandoc_paragraph_children(
            paragraph,
            element,
            hyperlink_map=hyperlink_map,
        )
        return True

    return False


def add_inline_content(paragraph, text):
    """
    追加段落内容：
    - 含公式：优先用 Pandoc 生成可编辑公式；
    - 不含公式：用本地行内样式解析器。
    """
    normalized_text = _simplify_link_label_line(text)

    if "$" in normalized_text and _append_pandoc_inline_to_paragraph(paragraph, normalized_text):
        return

    add_markdown_inline_runs(paragraph, normalized_text)

def preprocess_formula(text_content):
    """
    预处理公式，移除 Pandoc 不支持的命令
    """
    # 移除 \tag{} 命令，保留编号文本
    text_content = re.sub(r'\\tag\s*\{([^}]+)\}', r'(\1)', text_content)

    # 将 \left\{ \begin{array}{...} ... \end{array} \right. 转换为 \begin{cases} ... \end{cases}
    # 这是分段函数的两种不同写法，Pandoc 只支持 cases
    # 步骤1: 匹配并替换整个结构
    pattern = r'\\left\s*\\\{\s*\\begin\s*\{array\}\s*\{[^}]+\}(.*?)\\end\s*\{array\}\s*\\right\s*\.'

    def replace_with_cases(match):
        inner_content = match.group(1)
        return r'\begin{cases}' + inner_content + r'\end{cases}'

    text_content = re.sub(pattern, replace_with_cases, text_content, flags=re.DOTALL)

    return text_content

def convert_text_with_pandoc(text_content):
    """
    使用 Pandoc 转换包含公式的文本，返回 XML 元素列表
    """
    # 预处理公式
    text_content = preprocess_formula(text_content)

    # 创建临时 Markdown 文件
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp_md:
        tmp_md.write(text_content)
        tmp_md_path = tmp_md.name

    # 创建临时 docx 文件
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
        tmp_docx_path = tmp_docx.name

    # 使用 Pandoc 转换
    cmd = [
        'pandoc',
        tmp_md_path,
        '-o', tmp_docx_path,
        '-f', 'markdown+tex_math_dollars+strikeout+subscript+superscript+mark+autolink_bare_uris',
        '--standalone'
    ]

    elements = []
    hyperlink_map = {}
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True)

        # 读取生成的文档
        temp_doc = Document(tmp_docx_path)
        for rel in temp_doc.part.rels.values():
            if rel.reltype.endswith("/hyperlink"):
                hyperlink_map[rel.rId] = rel.target_ref

        # 提取所有 body 元素
        for element in temp_doc.element.body:
            elements.append(element)

    except subprocess.CalledProcessError as e:
        print(f"  警告：Pandoc 转换失败: {e}")
    finally:
        # 清理临时文件
        if os.path.exists(tmp_md_path):
            os.remove(tmp_md_path)
        if os.path.exists(tmp_docx_path):
            os.remove(tmp_docx_path)

    return {
        "elements": elements,
        "hyperlinks": hyperlink_map,
    }

def extract_table_caption(text_before_table):
    """
    从表格前的文本中提取表格标题（最后一个非空行）
    """
    lines = text_before_table.strip().split('\n')
    # 倒序查找，找到第一个非空行
    for i in range(len(lines) - 1, -1, -1):
        line = lines[i].strip()
        if line:
            # 检查是否是表格标题（TABLE 开头）
            if re.match(r'^TABLE\s*[IVX]+', line, re.IGNORECASE):
                # 返回标题和剩余文本
                caption = line
                remaining = '\n'.join(lines[:i]).strip()
                return caption, remaining
    return None, text_before_table

def parse_markdown_table(md_table_text):
    """
    解析 Markdown 表格
    返回：行列表，每行是单元格列表
    """
    lines = md_table_text.strip().split('\n')
    if len(lines) < 2:
        return None

    rows = []
    for i, line in enumerate(lines):
        # 跳过分隔行（如 | ---- | ---- |）
        if i == 1 and re.match(r'^\|[\s\-:|]+\|$', line):
            continue

        # 解析单元格
        cells = [cell.strip() for cell in line.split('|')]
        # 移除首尾空元素
        cells = [c for c in cells if c]

        if cells:
            rows.append(cells)

    return rows if rows else None

def create_word_table_from_markdown(doc, md_table_text, style_options=None):
    """
    从 Markdown 表格创建 Word 表格
    """
    rows_data = parse_markdown_table(md_table_text)
    if not rows_data or len(rows_data) == 0:
        return None

    num_rows = len(rows_data)
    num_cols = len(rows_data[0])

    # 创建表格
    table = doc.add_table(rows=num_rows, cols=num_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass

    # 填充内容
    for row_idx, row_cells in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_cells):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                add_inline_content(paragraph, cell_text)
                _apply_east_asia_font(paragraph, style_options=style_options)

    _apply_table_layout(table, apply_run_size=True)

    return table

def split_content_with_table_captions(content):
    """
    分割内容，识别表格及其标题（支持 HTML 表格和 Markdown 表格）
    返回：[(type, content, caption), ...]
    """
    parts = []

    # 先处理 HTML 表格
    html_table_pattern = r'<table>.*?</table>'
    # 再处理 Markdown 表格（连续的以 | 开头的行）
    md_table_pattern = r'(?:^\|.+\|$\n?)+'

    # 收集所有表格的位置（HTML 和 Markdown）
    all_tables = []

    # 查找 HTML 表格
    for match in re.finditer(html_table_pattern, content, re.DOTALL):
        all_tables.append(('html', match.start(), match.end(), match.group(0)))

    # 查找 Markdown 表格
    for match in re.finditer(md_table_pattern, content, re.MULTILINE):
        all_tables.append(('markdown', match.start(), match.end(), match.group(0)))

    # 按位置排序
    all_tables.sort(key=lambda x: x[1])

    last_pos = 0
    for table_type, start, end, table_content in all_tables:
        # 提取表格前的文本
        text_before = content[last_pos:start]

        # 尝试从文本中提取表格标题
        caption, remaining_text = extract_table_caption(text_before)

        # 如果有剩余文本，添加为文本块
        if remaining_text.strip():
            parts.append(('text', remaining_text.strip(), None))

        # 添加表格（带标题和类型）
        parts.append((f'table_{table_type}', table_content, caption))
        last_pos = end

    # 添加最后的文本
    if last_pos < len(content):
        text_after = content[last_pos:].strip()
        if text_after:
            parts.append(('text', text_after, None))

    return parts

def split_text_into_segments(text_content):
    """
    将文本内容分割成不同的片段
    """
    segments = []

    # 首先提取多行公式块 $$...$$
    formula_block_pattern = r'\$\$[\s\S]*?\$\$'
    parts = []
    last_pos = 0

    for match in re.finditer(formula_block_pattern, text_content):
        # 添加公式块前的文本
        if match.start() > last_pos:
            parts.append(('text', text_content[last_pos:match.start()]))

        # 添加公式块
        parts.append(('formula_block', match.group(0)))
        last_pos = match.end()

    # 添加最后的文本
    if last_pos < len(text_content):
        parts.append(('text', text_content[last_pos:]))

    # 进一步处理文本部分（逐行）
    for part_type, part_content in parts:
        if part_type == 'formula_block':
            segments.append(('formula_block', part_content))
        else:
            lines = part_content.split('\n')
            idx = 0

            while idx < len(lines):
                line = lines[idx]
                line_stripped = line.strip()
                if not line_stripped:
                    idx += 1
                    continue

                fence_match = FENCE_START_RE.match(line_stripped)
                if fence_match:
                    fence = fence_match.group("fence")
                    idx += 1
                    code_lines = []

                    while idx < len(lines):
                        close_line = lines[idx].strip()
                        if close_line.startswith(fence):
                            idx += 1
                            break
                        code_lines.append(lines[idx])
                        idx += 1

                    segments.append(("code_block", "\n".join(code_lines)))
                    continue

                # 列表块（支持多级嵌套，保留缩进原文）
                if LIST_ITEM_RE.match(line):
                    list_lines = [line]
                    idx += 1

                    while idx < len(lines):
                        next_line = lines[idx]
                        next_stripped = next_line.strip()

                        if not next_stripped:
                            list_lines.append(next_line)
                            idx += 1
                            continue

                        if LIST_ITEM_RE.match(next_line):
                            list_lines.append(next_line)
                            idx += 1
                            continue

                        # 列表项续行/嵌套内容（缩进行）
                        if next_line.startswith(' ') or next_line.startswith('\t'):
                            list_lines.append(next_line)
                            idx += 1
                            continue

                        break

                    while list_lines and not list_lines[-1].strip():
                        list_lines.pop()

                    if list_lines:
                        segments.append(('list_block', '\n'.join(list_lines)))
                    continue

                # 检测图片
                if HORIZONTAL_RULE_RE.match(line_stripped):
                    segments.append(('horizontal_rule', line_stripped))
                # 检测图片
                elif re.match(r'!\[.*?\]\(.*?\)', line_stripped):
                    segments.append(('image', line_stripped))
                # 检测标题
                elif line_stripped.startswith('#'):
                    segments.append(('heading', line_stripped))
                # 检测单行公式（包含 $）
                elif '$' in line_stripped:
                    segments.append(('formula_line', line_stripped))
                # 普通文本
                else:
                    segments.append(('text', line_stripped))

                idx += 1

    return segments


def _indent_to_level(indent_text):
    """按缩进宽度估算列表层级（2 空格为一层，tab 按 4 空格）。"""
    width = 0
    for ch in indent_text:
        width += 4 if ch == '\t' else 1
    return max(width // 2, 0)


def _resolve_list_style_name(doc, is_ordered, level):
    base = 'List Number' if is_ordered else 'List Bullet'
    capped_level = min(level + 1, 3)
    candidates = [f'{base} {capped_level}', base] if level > 0 else [base]

    for name in candidates:
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue

    return None


def process_list_block(doc, list_content, render_state=None):
    """将 Markdown 列表块转换为 Word 列表段落。"""
    current_para = None
    style_options = (render_state or {}).get("style_options")

    for raw_line in list_content.split('\n'):
        if not raw_line.strip():
            current_para = None
            continue

        match = LIST_ITEM_RE.match(raw_line)
        if match:
            indent = match.group('indent')
            marker = match.group('marker')
            item_text = match.group('content')

            level = _indent_to_level(indent)
            is_ordered = marker[0].isdigit()

            para = doc.add_paragraph()
            style_name = _resolve_list_style_name(doc, is_ordered, level)
            if style_name:
                para.style = style_name

            # 超过内置 3 级样式时，额外缩进保证层级可读
            if level >= 3:
                para.paragraph_format.left_indent = Inches(0.25 * level)

            add_inline_content(para, item_text)
            _apply_east_asia_font(
                para,
                style_options=style_options,
                apply_spacing=True,
            )
            current_para = para
            continue

        continuation_text = raw_line.strip()
        if not continuation_text:
            continue

        if current_para is None:
            para = doc.add_paragraph()
            add_inline_content(para, continuation_text)
            _apply_east_asia_font(
                para,
                style_options=style_options,
                apply_spacing=True,
            )
            current_para = para
            continue

        if current_para.runs:
            current_para.add_run(' ')
        add_inline_content(current_para, continuation_text)
        _apply_east_asia_font(
            current_para,
            style_options=style_options,
            apply_spacing=True,
        )

def process_segment(doc, seg_type, seg_content, base_dir=None, render_state=None):
    """
    处理单个片段
    base_dir: markdown 文件所在目录，用于解析相对图片路径
    """
    style_options = (render_state or {}).get("style_options")

    if seg_type == 'image':
        # 处理图片
        image_pattern = r'!\[(.*?)\]\((.*?)\)'
        match = re.match(image_pattern, seg_content)
        if match:
            alt_text = match.group(1)
            image_path = match.group(2)

            # 如果是相对路径，基于 md 文件目录解析
            if base_dir and not os.path.isabs(image_path):
                image_path = os.path.join(base_dir, image_path)

            if os.path.exists(image_path):
                try:
                    print(f"  - 添加图片: {image_path}")
                    # 创建段落并设置居中对齐
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 在居中的段落中添加图片
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(4.5))

                    # 添加图片说明（居中）
                    if alt_text:
                        para = doc.add_paragraph(alt_text)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _apply_east_asia_font(
                            para,
                            style_options=style_options,
                            apply_spacing=True,
                        )
                except Exception as e:
                    print(f"  警告：添加图片失败 {image_path}: {e}")
                    doc.add_paragraph(f"[图片: {image_path}]")
            else:
                print(f"  警告：图片不存在: {image_path}")
                doc.add_paragraph(f"[图片不存在: {image_path}]")

    elif seg_type == 'heading':
        # 处理标题
        level = len(re.match(r'^#+', seg_content).group())
        title_text = seg_content.lstrip('#').strip()
        _update_chapter_state_by_heading(level, render_state)
        heading = doc.add_heading(level=min(level, 9))
        add_inline_content(heading, title_text)
        # 显式设置标题字体
        _apply_east_asia_font(
            heading,
            style_options=style_options,
            apply_spacing=True,
        )

    elif seg_type == 'formula_block' or seg_type == 'formula_line':
        # 处理公式（用 Pandoc）
        print(f"  - 处理公式")
        payload = convert_text_with_pandoc(seg_content)
        elements = payload["elements"]
        hyperlink_map = payload["hyperlinks"]
        numbering_mode = (render_state or {}).get(
            "formula_numbering_mode",
            FORMULA_NUMBERING_CHAPTER_INDEX,
        )
        if (
            seg_type == 'formula_block'
            and render_state
            and numbering_mode != FORMULA_NUMBERING_NONE
        ):
            label = _next_block_formula_label(render_state)
            _append_numbered_formula_block(
                doc,
                elements,
                label,
                hyperlink_map=hyperlink_map,
                style_options=style_options,
            )
        else:
            for element in _extract_pandoc_paragraphs(elements):
                # 创建一个新段落，然后复制 Pandoc 生成的段落内容
                new_para = doc.add_paragraph()
                _copy_pandoc_paragraph_children(
                    new_para,
                    element,
                    hyperlink_map=hyperlink_map,
                )
                # 设置段落中所有文本的字体
                _apply_east_asia_font(
                    new_para,
                    style_options=style_options,
                    apply_spacing=True,
                )

    elif seg_type == 'list_block':
        process_list_block(doc, seg_content, render_state=render_state)

    elif seg_type == 'horizontal_rule':
        # 水平分割线
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        para = doc.add_paragraph()
        p_pr = para._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    elif seg_type == 'code_block':
        # 代码块
        from docx.oxml.shared import qn

        para = doc.add_paragraph()
        run = para.add_run(seg_content)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        run_pr = run._element.get_or_add_rPr()
        run_pr.rFonts.set(qn('w:ascii'), 'Consolas')
        run_pr.rFonts.set(qn('w:hAnsi'), 'Consolas')
        run_pr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        _apply_paragraph_spacing(para, style_options=style_options)

    elif seg_type == 'text':
        # 普通文本
        para = doc.add_paragraph()
        add_inline_content(para, seg_content)

        # 显式设置段落字体
        _apply_east_asia_font(
            para,
            style_options=style_options,
            apply_spacing=True,
        )

        # 检查是否是图片说明（通常以括号开头或以 Fig. 开头）
        # 例如: "(a) Description..." 或 "Fig. 1. Description..."
        # 但排除列表项如 "(1) 长文本..."
        stripped = seg_content.strip()
        # 只有当括号内是单个字母或短标识符时才认为是图片说明
        is_figure_caption = (
            stripped.startswith('Fig.') or
            stripped.startswith('Figure') or
            re.match(r'^图\s*\d', stripped) or  # 中文图标题：图1、图 3.1 等
            (stripped.startswith('(') and re.match(r'^\([a-zA-Z]\)\s', stripped))  # 只匹配 (a) (b) 等
        )
        if is_figure_caption:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_text_block(doc, text_content, base_dir=None, render_state=None):
    """
    处理文本块：分割成不同类型的片段并分别处理
    """
    segments = split_text_into_segments(text_content)

    for seg_type, seg_content in segments:
        process_segment(
            doc,
            seg_type,
            seg_content,
            base_dir=base_dir,
            render_state=render_state,
        )

def add_table_caption_with_formula(doc, caption_text, render_state=None):
    """
    添加表格标题（可能包含公式）
    """
    if '$' in caption_text:
        # 标题包含公式，用 Pandoc 处理
        print(f"  - 添加表格标题（含公式）")
        payload = convert_text_with_pandoc(caption_text)
        elements = payload["elements"]
        hyperlink_map = payload["hyperlinks"]

        # 找到所有段落元素
        for element in elements:
            # 只添加段落元素，跳过 sectPr（分节符）等格式元素
            if element.tag.endswith('}p'):
                # 创建一个新段落，然后复制 Pandoc 生成的段落内容
                new_para = doc.add_paragraph()

                # 复制所有子元素（包括公式）到新段落
                _copy_pandoc_paragraph_children(
                    new_para,
                    element,
                    hyperlink_map=hyperlink_map,
                )

                # 设置段落居中对齐
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _apply_east_asia_font(
                    new_para,
                    style_options=(render_state or {}).get("style_options"),
                    apply_spacing=True,
                )
    else:
        # 普通标题
        para = doc.add_paragraph()
        add_inline_content(para, caption_text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.style = 'Caption'
        _apply_east_asia_font(
            para,
            style_options=(render_state or {}).get("style_options"),
            apply_spacing=True,
        )

def convert_with_python_docx(
    input_file,
    output_file,
    enable_block_formula_numbering=True,
    formula_numbering_mode=None,
    doc_style_options=None,
):
    """
    完全使用 python-docx 创建文档（表格 + 图片 + 公式）
    """
    if not os.path.exists(input_file):
        print(f"错误：找不到文件 {input_file}")
        return

    print(f"正在读取文件: {input_file}")
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 图片路径基于 md 文件所在目录解析
    base_dir = os.path.dirname(os.path.abspath(input_file))

    # 创建文档
    doc = Document()

    style_options = normalize_doc_style_options(doc_style_options)
    numbering_mode = normalize_formula_numbering_mode(
        formula_numbering_mode,
        enable_block_formula_numbering=enable_block_formula_numbering,
    )

    # 设置默认字体
    from docx.oxml.shared import qn

    # 设置正文字体
    style = doc.styles['Normal']
    font = style.font
    font.name = style_options["font_en"]
    font.size = Pt(style_options["font_size_body"])
    style_rpr = style._element.get_or_add_rPr()
    style_rpr.rFonts.set(qn('w:ascii'), style_options["font_en"])
    style_rpr.rFonts.set(qn('w:hAnsi'), style_options["font_en"])
    style_rpr.rFonts.set(qn('w:eastAsia'), style_options["font_zh"])
    style.paragraph_format.line_spacing = style_options["line_spacing"]
    style.paragraph_format.space_after = Pt(style_options["paragraph_spacing"])

    # 设置所有标题样式的字体
    for i in range(1, 10):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = style_options["font_en"]
            if i == 1:
                heading_font.size = Pt(style_options["font_size_h1"])
            elif i == 2:
                heading_font.size = Pt(style_options["font_size_h2"])
            elif i == 3:
                heading_font.size = Pt(style_options["font_size_h3"])
            elif i == 4:
                heading_font.size = Pt(style_options["font_size_h4"])
            heading_rpr = heading_style._element.get_or_add_rPr()
            heading_rpr.rFonts.set(qn('w:ascii'), style_options["font_en"])
            heading_rpr.rFonts.set(qn('w:hAnsi'), style_options["font_en"])
            heading_rpr.rFonts.set(qn('w:eastAsia'), style_options["font_zh"])
            heading_style.paragraph_format.line_spacing = style_options["line_spacing"]
            heading_style.paragraph_format.space_after = Pt(style_options["paragraph_spacing"])
        except KeyError:
            pass  # 如果样式不存在，跳过

    render_state = {
        'formula_numbering_mode': numbering_mode,
        'chapter_level': None,
        'current_chapter': 0,
        'global_formula_counter': 0,
        'chapter_formula_counter': 0,
        'style_options': style_options,
    }

    # 分割内容（识别表格标题）
    print("正在解析文档内容...")
    parts = split_content_with_table_captions(content)

    # 处理每个部分
    print("正在构建 Word 文档...")
    for part_type, part_content, caption in parts:
        if part_type == 'table_html':
            # HTML 表格
            # 如果有标题，先添加标题
            if caption:
                add_table_caption_with_formula(doc, caption, render_state=render_state)

            # 添加表格
            print("  - 添加 HTML 表格")
            create_word_table(doc, part_content, style_options=style_options)
            spacer = doc.add_paragraph()  # 表格后添加空行
            _apply_paragraph_spacing(spacer, style_options=style_options)
        elif part_type == 'table_markdown':
            # Markdown 表格
            # 如果有标题，先添加标题
            if caption:
                add_table_caption_with_formula(doc, caption, render_state=render_state)

            # 添加表格
            print("  - 添加 Markdown 表格")
            create_word_table_from_markdown(doc, part_content, style_options=style_options)
            spacer = doc.add_paragraph()  # 表格后添加空行
            _apply_paragraph_spacing(spacer, style_options=style_options)
        else:
            # 处理文本块（包括标题、图片、公式、段落）
            process_text_block(
                doc,
                part_content,
                base_dir=base_dir,
                render_state=render_state,
            )

    # 保存文档
    doc.save(output_file)
    print(f"\n转换成功！已保存到: {output_file}")

if __name__ == "__main__":
    import sys

    # 支持命令行参数
    if len(sys.argv) > 1:
        input_md = sys.argv[1]
        output_docx = sys.argv[2] if len(sys.argv) > 2 else input_md.replace('.md', '.docx')
    else:
        # 默认转换 paper.md
        input_md = 'paper.md'
        output_docx = 'paper.docx'

    convert_with_python_docx(input_md, output_docx)
